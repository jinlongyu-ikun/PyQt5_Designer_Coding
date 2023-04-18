# author jinyunlong
# createtime 2023/4/18 9:56
# 职业 锅炉房保安

import sys
from PyQt5.QtWidgets import (QApplication, QMainWindow, QTextEdit, QPushButton, QVBoxLayout, QHBoxLayout, QWidget,
                             QFileDialog, QLabel, QLineEdit, QListWidget, QSizePolicy)
from PyQt5.QtCore import QThread
from docx import Document
import pyttsx3


class SpeechThread(QThread):
    def __init__(self, text, parent=None):
        super(SpeechThread, self).__init__(parent)
        self.text = text
        self.tts_engine = pyttsx3.init()

    def run(self):
        self.tts_engine.say(self.text)
        self.tts_engine.runAndWait()


class DocumentVisualizer(QMainWindow):
    def __init__(self):
        super().__init__()

        self.text_display = QTextEdit(self)
        self.text_display.setReadOnly(True)

        self.open_button = QPushButton('Open', self)
        self.open_button.clicked.connect(self.open_document)

        self.search_input = QLineEdit(self)
        self.search_button = QPushButton('Search', self)
        self.search_button.clicked.connect(self.search_document)

        self.paragraphs_list = QListWidget(self)
        self.paragraphs_list.itemClicked.connect(self.go_to_paragraph)

        self.restore_button = QPushButton('Restore', self)
        self.restore_button.clicked.connect(self.restore_document)

        layout = QVBoxLayout()

        top_bar = QHBoxLayout()
        top_bar.addWidget(self.open_button)
        top_bar.addWidget(QLabel('Search:'))
        top_bar.addWidget(self.search_input)
        top_bar.addWidget(self.search_button)
        layout.addLayout(top_bar)

        middle_section = QHBoxLayout()
        middle_section.addWidget(self.text_display)
        middle_section.addWidget(self.paragraphs_list)
        layout.addLayout(middle_section)

        layout.addWidget(self.restore_button)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        self.setWindowTitle('Document Visualizer')
        self.setGeometry(100, 100, 800, 600)

        self.document = None

        self.setStyleSheet("""
            QMainWindow {
                background-color: #f0f0f0;
            }
            QPushButton {
                background-color: #0099cc;
                color: white;
                padding: 5px 15px;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #0073a8;
            }
            QTextEdit {
                background-color: #ffffff;
                border: 1px solid #ccc;
                border-radius: 5px;
            }
            QLineEdit {
                background-color: #ffffff;
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 2px 5px;
            }
            QListWidget {
                background-color: #ffffff;
                border: 1px solid #ccc;
                border-radius: 5px;
            }
        """)

    def open_document(self):
        file_name, _ = QFileDialog.getOpenFileName(self, 'Open Document', '', 'Word Documents (*.docx)')

        if file_name:
            self.document = Document(file_name)
            text = '\n'.join(p.text for p in self.document.paragraphs)
            self.text_display.setPlainText(text)

            self.paragraphs_list.clear()
            for i, p in enumerate(self.document.paragraphs):
                self.paragraphs_list.addItem(f"Paragraph {i}: {p.text[:30]}...")

    def search_document(self):
        if not self.document:
            return

        query = self.search_input.text()

        # Check if the query is a number
        if query.isdigit():
            paragraph_index = int(query)
            if 0 <= paragraph_index < len(self.document.paragraphs):
                paragraph_text = self.document.paragraphs[paragraph_index].text
                self.text_display.setPlainText(f"Paragraph {paragraph_index}: {paragraph_text}")
                self.update_table_of_contents([(paragraph_index, self.document.paragraphs[paragraph_index])])
            else:
                self.text_display.setPlainText('No results found.')
                self.update_table_of_contents()
        else:
            results = []

            for i, p in enumerate(self.document.paragraphs):
                if query in p.text:
                    results.append((i, p))

            if not results:
                self.text_display.setPlainText('No results found.')
                self.update_table_of_contents()
            else:
                result_text = '\n'.join(f"Paragraph {i}: {p.text}" for i, p in results)
                self.text_display.setPlainText(result_text)
                self.update_table_of_contents(results)

    def go_to_paragraph(self, item):
        item_text = item.text()
        paragraph_index = int(item_text.split(":")[0].split(" ")[1])
        paragraph_text = self.document.paragraphs[paragraph_index].text
        self.text_display.setPlainText(paragraph_text)

        if hasattr(self, 'speech_thread'):
            self.speech_thread.terminate()

        if paragraph_text.strip() != "...":
            self.speech_thread = SpeechThread(paragraph_text)
            self.speech_thread.start()

    # Add this function to the DocumentVisualizer class
    def restore_document(self):
        if self.document:
            text = '\n'.join(p.text for p in self.document.paragraphs)
            self.text_display.setPlainText(text)
            self.update_table_of_contents()

    def update_table_of_contents(self, results=None):
        self.paragraphs_list.clear()
        if results is None:
            results = enumerate(self.document.paragraphs)

        for i, p in results:
            self.paragraphs_list.addItem(f"Paragraph {i}: {p.text[:30]}...")


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = DocumentVisualizer()
    window.show()
    sys.exit(app.exec_())

